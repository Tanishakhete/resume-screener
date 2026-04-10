"""
Data Preprocessing Module for Resume Screening System.
Handles text extraction, cleaning, and preprocessing.
"""

import re
import string
import nltk
import PyPDF2
import docx
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize
from nltk.stem import WordNetLemmatizer


# Download required NLTK data
def download_nltk_data():
    """Download required NLTK data packages."""
    try:
        nltk.data.find('tokenizers/punkt')
    except LookupError:
        nltk.download('punkt', quiet=True)
    
    try:
        nltk.data.find('corpora/stopwords')
    except LookupError:
        nltk.download('stopwords', quiet=True)
    
    try:
        nltk.data.find('corpora/wordnet')
    except LookupError:
        nltk.download('wordnet', quiet=True)
    
    try:
        nltk.data.find('taggers/averaged_perceptron_tagger')
    except LookupError:
        nltk.download('averaged_perceptron_tagger', quiet=True)


class TextExtractor:
    """Extract text from PDF and DOCX files."""
    
    @staticmethod
    def extract_from_pdf(file_path):
        """
        Extract text from PDF file.
        
        Args:
            file_path: Path to PDF file or file-like object
            
        Returns:
            str: Extracted text
        """
        text = ""
        try:
            if hasattr(file_path, 'read'):
                pdf_reader = PyPDF2.PdfReader(file_path)
            else:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
            
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        except Exception as e:
            print(f"Error extracting PDF: {e}")
        return text
    
    @staticmethod
    def extract_from_docx(file_path):
        """
        Extract text from DOCX file.
        
        Args:
            file_path: Path to DOCX file or file-like object
            
        Returns:
            str: Extracted text
        """
        text = ""
        try:
            if hasattr(file_path, 'read'):
                doc = docx.Document(file_path)
            else:
                doc = docx.Document(file_path)
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
        except Exception as e:
            print(f"Error extracting DOCX: {e}")
        return text
    
    @staticmethod
    def extract_text(file_path, file_type=None):
        """
        Extract text from file based on type.
        
        Args:
            file_path: Path to file or file-like object
            file_type: File extension (pdf, docx) or None to auto-detect
            
        Returns:
            str: Extracted text
        """
        if file_type is None:
            if hasattr(file_path, 'name'):
                file_type = file_path.name.split('.')[-1].lower()
            elif isinstance(file_path, str):
                file_type = file_path.split('.')[-1].lower()
        
        if file_type == 'pdf':
            return TextExtractor.extract_from_pdf(file_path)
        elif file_type in ['docx', 'doc']:
            return TextExtractor.extract_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")


class TextPreprocessor:
    """Preprocess text data for machine learning."""
    
    def __init__(self):
        """Initialize preprocessor with required tools."""
        download_nltk_data()
        self.lemmatizer = WordNetLemmatizer()
        self.stop_words = set(stopwords.words('english'))
        
        # Add custom stopwords relevant to resumes
        custom_stopwords = {
            'resume', 'cv', 'curriculum', 'vitae', 'page', 'email', 'phone',
            'contact', 'address', 'linkedin', 'github', 'portfolio', 'name',
            'experience', 'education', 'skills', 'objective', 'summary'
        }
        self.stop_words.update(custom_stopwords)
    
    def clean_text(self, text):
        """
        Clean raw text by removing special characters, URLs, emails, etc.
        
        Args:
            text: Raw text string
            
        Returns:
            str: Cleaned text
        """
        # Convert to lowercase
        text = text.lower()
        
        # Remove URLs
        text = re.sub(r'http\S+|www\S+|https\S+', '', text, flags=re.MULTILINE)
        
        # Remove email addresses
        text = re.sub(r'\S+@\S+', '', text)
        
        # Remove phone numbers
        text = re.sub(r'\b\d{3}[-.]?\d{3}[-.]?\d{4}\b', '', text)
        
        # Remove special characters and digits, keep only letters and spaces
        text = re.sub(r'[^a-zA-Z\s]', ' ', text)
        
        # Remove extra whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove single characters
        text = re.sub(r'\b[a-zA-Z]\b', '', text)
        
        return text.strip()
    
    def tokenize(self, text):
        """
        Tokenize text into words.
        
        Args:
            text: Text string
            
        Returns:
            list: List of tokens
        """
        return word_tokenize(text)
    
    def remove_stopwords(self, tokens):
        """
        Remove stopwords from token list.
        
        Args:
            tokens: List of tokens
            
        Returns:
            list: Filtered tokens
        """
        return [token for token in tokens if token not in self.stop_words and len(token) > 2]
    
    def lemmatize(self, tokens):
        """
        Lemmatize tokens to their base form.
        
        Args:
            tokens: List of tokens
            
        Returns:
            list: Lemmatized tokens
        """
        return [self.lemmatizer.lemmatize(token) for token in tokens]
    
    def preprocess(self, text):
        """
        Complete preprocessing pipeline.
        
        Args:
            text: Raw text string
            
        Returns:
            str: Preprocessed text
        """
        # Clean text
        cleaned = self.clean_text(text)
        
        # Tokenize
        tokens = self.tokenize(cleaned)
        
        # Remove stopwords
        tokens = self.remove_stopwords(tokens)
        
        # Lemmatize
        tokens = self.lemmatize(tokens)
        
        # Join tokens back to string
        return ' '.join(tokens)
    
    def preprocess_resume(self, file_path, file_type=None):
        """
        Complete pipeline: extract and preprocess resume.
        
        Args:
            file_path: Path to resume file
            file_type: File type (pdf, docx)
            
        Returns:
            dict: Dictionary with raw and processed text
        """
        # Extract text
        raw_text = TextExtractor.extract_text(file_path, file_type)
        
        # Preprocess
        processed_text = self.preprocess(raw_text)
        
        return {
            'raw_text': raw_text,
            'processed_text': processed_text
        }


def load_and_preprocess_dataset(file_path, text_column, label_column):
    """
    Load and preprocess a dataset from CSV file.
    
    Args:
        file_path: Path to CSV file
        text_column: Name of the text column
        label_column: Name of the label column
        
    Returns:
        tuple: (preprocessed_texts, labels, preprocessor)
    """
    import pandas as pd
    
    # Load dataset
    df = pd.read_csv(file_path)
    
    # Initialize preprocessor
    preprocessor = TextPreprocessor()
    
    # Preprocess all texts
    preprocessed_texts = []
    for text in df[text_column]:
        if pd.isna(text):
            preprocessed_texts.append('')
        else:
            preprocessed_texts.append(preprocessor.preprocess(str(text)))
    
    labels = df[label_column].values
    
    return preprocessed_texts, labels, preprocessor


if __name__ == "__main__":
    # Test the preprocessing module
    print("Testing Data Preprocessing Module...")
    
    # Test text preprocessor
    sample_text = """
    John Doe
    Email: john.doe@email.com | Phone: 123-456-7890
    LinkedIn: linkedin.com/in/johndoe
    
    EXPERIENCE
    Software Engineer at ABC Corp (2020-Present)
    - Developed machine learning models using Python and TensorFlow
    - Implemented data pipelines for processing large datasets
    
    EDUCATION
    Bachelor of Science in Computer Science, XYZ University
    
    SKILLS
    Python, Machine Learning, Data Science, SQL, TensorFlow
    """
    
    preprocessor = TextPreprocessor()
    processed = preprocessor.preprocess(sample_text)
    
    print("\nOriginal text length:", len(sample_text))
    print("Processed text length:", len(processed))
    print("Processed text:", processed[:200] + "...")
